# app.py
import streamlit as st
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_chroma import Chroma
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.runnables.history import RunnableWithMessageHistory
from dotenv import load_dotenv
import os
import fitz  # PyMuPDF
from docx import Document
import json

# Import database functions
from database import save_document, get_documents
from link_processing import process_link  # Import the link processing functions

# Load environment variables
load_dotenv()

# Constants
API_KEY = os.getenv("GROQ_API_KEY")
HF_TOKEN = os.getenv("HF_TOKEN")

# Set up embeddings and LLM
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
llm = ChatGroq(groq_api_key=API_KEY, model_name="Gemma2-9b-It")

# Set up Streamlit layout
st.set_page_config(layout="wide")

# Initialize documents
documents = []

# Document wrapper class
class SimpleDocument:
    def __init__(self, content: str, metadata: dict = None):
        self.page_content = content
        self.metadata = metadata or {}

# Define session history functions before usage
if 'store' not in st.session_state:
    st.session_state.store = {}
    st.session_state.documents_processed = False  # Flag to check if documents have been processed
    st.session_state.messages = []  # Initialize messages

def get_session_history(session: str) -> BaseChatMessageHistory:
    if session not in st.session_state.store:
        st.session_state.store[session] = ChatMessageHistory()
    return st.session_state.store[session]

def truncate_messages(messages, max_length=4000):
    total_length = 0
    truncated_messages = []
    
    for message in reversed(messages):
        message_length = len(message['content'])
        if total_length + message_length > max_length:
            break
        truncated_messages.append(message)
        total_length += message_length
    
    return list(reversed(truncated_messages))

# Sidebar: File Uploads, Link Input, and Clear Chat Button
with st.sidebar:
    st.header("Upload Documents or Enter Links")
    
    # File uploader
    uploaded_files = st.file_uploader(
        "Choose files (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True
    )
    
    # Link input
    link_input = st.text_input("Or enter a link (YouTube, website):", key="link_input", label_visibility='hidden', placeholder="Link website/Youtube")

    if st.button("Clear Chat"):
        st.session_state.messages = []
    
    # Handle file uploads
    def process_pdf(file):
        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
        text = "".join(page.get_text() for page in pdf_document)
        return text

    def process_docx(file):
        doc = Document(file)
        text = "\n".join(para.text for para in doc.paragraphs)
        return text

    def process_txt(file):
        text = file.read().decode("utf-8")
        return text

    if uploaded_files:
        documents = []
        for uploaded_file in uploaded_files:
            if uploaded_file.type == "application/pdf":
                text = process_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = process_docx(uploaded_file)
            elif uploaded_file.type == "text/plain":
                text = process_txt(uploaded_file)
            else:
                st.warning("Unsupported file type.")
                continue
            
            # Wrap text in SimpleDocument with metadata
            metadata = {"source": uploaded_file.name}
            documents.append(SimpleDocument(content=text, metadata=metadata))

        # Save documents to the database
        for doc in documents:
            save_document(doc.page_content, doc.metadata)

    if link_input:
        text = process_link(link_input)
        if text != "Unsupported link type.":
            documents.append(SimpleDocument(content=text, metadata={"source": link_input}))
            save_document(text, {"source": link_input})
        else:
            st.warning("Unsupported link type.")
    
    # Retrieve documents and create vectorstore
    saved_documents = get_documents()
    documents = [SimpleDocument(content=doc[1], metadata=json.loads(doc[2])) for doc in saved_documents]
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=500)
    splits = text_splitter.split_documents(documents)

    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
    retriever = vectorstore.as_retriever()

    # Create chat and QA chains
    contextualize_q_system_prompt = (
        "Given a chat history and the latest user question, "
        "which might reference context in the chat history, "
        "reformulate the question if needed. Do not answer the question."
    )
    contextualize_q_prompt = ChatPromptTemplate.from_messages(
        [("system", contextualize_q_system_prompt), MessagesPlaceholder("chat_history"), ("human", "{input}")]
    )
    history_aware_retriever = create_history_aware_retriever(llm, retriever, contextualize_q_prompt)

    system_prompt = (
        "You are an assistant for question-answering tasks. "
        "Use the following pieces of retrieved context to answer "
        "the question. If you don't know the answer, say that you "
        "don't know. Use three sentences maximum."
        "\n\n{context}"
    )
    qa_prompt = ChatPromptTemplate.from_messages(
        [("system", system_prompt), MessagesPlaceholder("chat_history"), ("human", "{input}")]
    )
    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

    conversational_rag_chain = RunnableWithMessageHistory(
        rag_chain, get_session_history,
        input_messages_key="input",
        history_messages_key="chat_history",
        output_messages_key="answer"
    )
    st.session_state.documents_processed = True  # Set flag to True

# Main area: Chat Interface
st.title("Chat with 'Doc'")
st.write("Upload Documents, Enter Links, And Ask Anything About Them")

# Display chat history using a placeholder
chat_placeholder = st.empty()

# Input field and Ask button
user_input = st.text_input("Your question:", key="user_input", label_visibility='hidden')
ask_button = st.button("Ask")

def handle_user_input(user_input):
    if not st.session_state.documents_processed:
        st.warning("No documents have been processed yet.")
        return
    
    session_id = "default_session"  # Use a fixed session ID
    session_history = get_session_history(session_id)

    # Truncate messages before invoking the model
    truncated_messages = truncate_messages(st.session_state.messages)

    response = conversational_rag_chain.invoke(
        {"input": user_input},
        config={"configurable": {"session_id": session_id}}
    )

    # Update session history with the new question and answer
    st.session_state.messages.append({'type': 'human', 'content': user_input})
    st.session_state.messages.append({'type': 'assistant', 'content': response['answer']})

    # Update chat history display
    with chat_placeholder.container():
        st.write(
            """
            <style>
                .chat-container {
                    display: flex;
                    flex-direction: column;
                    height: 500px;
                    overflow-y: auto;
                }
                .message {
                    margin: 10px;
                    padding: 10px;
                    border-radius: 10px;
                }
                .user-message {
                    background-color: black;
                }
                .assistant-message {
                    background-color: white;
                    color:black;
                }
                .user-message::before {
                    content: "You: ";
                    font-weight: bold;
                    color: white;
                }
                .assistant-message::before {
                    content: "Assistant: ";
                    font-weight: bold;
                    color: black;
                }
            </style>
            """,
            unsafe_allow_html=True,
        )
        chat_container = st.empty()

        with chat_container.container():
            for message in st.session_state.messages:
                message_class = "user-message" if message['type'] == 'human' else "assistant-message"
                st.markdown(f'<div class="message {message_class}">{message["content"]}</div>', unsafe_allow_html=True)

    # Ensure the chat container scrolls to the bottom
    st.write(
        """
        <script>
            var chatContainer = document.getElementsByClassName('chat-container')[0];
            chatContainer.scrollTop = chatContainer.scrollHeight;
        </script>
        """,
        unsafe_allow_html=True,
    )

if ask_button and user_input:
    handle_user_input(user_input)